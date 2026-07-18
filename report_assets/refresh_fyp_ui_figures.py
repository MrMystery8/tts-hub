#!/usr/bin/env python3
"""Refresh the 13 live UI figures and matching prose in the private FYP DOCX."""

from __future__ import annotations

import hashlib
import shutil
import zipfile
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Ayaan Minhas-TP077859-APD3F2511CS(AI)-FYP Final Report (DRAFT).docx"
IMAGES = ROOT / "report_assets" / "ui_screenshots"
BACKUPS = ROOT / "report_assets" / "backups"

FIGURES = {
    "4.12": "Figure_4.12_Desktop_Generate.png",
    "4.13": "Figure_4.13_Desktop_Voices.png",
    "4.14": "Figure_4.14_Desktop_Reference_Intake.png",
    "4.15": "Figure_4.15_Desktop_Model_Rail.png",
    "4.16": "Figure_4.16_Desktop_Jobs.png",
    "4.17": "Figure_4.17_Desktop_Watermark_Lab.png",
    "4.18": "Figure_4.18_Desktop_Tour_Navigation.png",
    "4.19": "Figure_4.19_Desktop_Tour_Reference.png",
    "4.20": "Figure_4.20_Desktop_Tour_Actions.png",
    "4.21": "Figure_4.21_Mobile_Generate.png",
    "4.22": "Figure_4.22_Mobile_Jobs.png",
    "4.23": "Figure_4.23_Mobile_Voices.png",
    "4.24": "Figure_4.24_Mobile_Verify.png",
}

CAPTIONS = {
    "4.12": "Figure 4.12: Quick Phrases, speech generation and save-enabled output player.",
    "4.16": "Figure 4.16: Generation results, labels, saved phrases and output history view.",
    "4.18": "Figure 4.18: Guided tour, step 1 of 9 - the navigation rail and the five surfaces of the hub.",
    "4.19": "Figure 4.19: Guided tour, step 5 of 9 - the reference voice panel, spotlighted.",
    "4.20": "Figure 4.20: Guided tour, step 7 of 9 - output format, provenance watermark and the generate control.",
    "4.21": "Figure 4.21: Mobile companion - Generate screen with Quick Phrases first.",
    "4.22": "Figure 4.22: Mobile companion - labelled Jobs with the in-place saved-audio player.",
}

PROSE = {
    "4.12": (
        "The Generate screen in Figure 4.12 brings the complete synthesis workflow into one view. "
        "A labelled Quick Phrase is pinned above the script editor and replays its already-rendered local audio without submitting another generation request. "
        "Qwen3-TTS MLX, the persistent FYP Demo voice, WAV output and provenance watermarking are visible in the same workspace. "
        "The primary action changes to Generating and becomes unavailable while a queued or active job exists, while cancellation remains visible. "
        "The persistent result dock supplies playback, waveform feedback, duration metadata, download access and a direct Save control for adding or removing the current clip from Quick Phrases. "
        "All text, reference audio, job metadata and generated speech remain on the user's machine."
    ),
    "4.16": (
        "Figure 4.16 shows the completed-job history produced by the three frozen backends using curated, neutral text. "
        "A star column and Saved filter support curation, while the labelled Qwen result opens a detail panel with playback, download, rename, restore and deletion controls. "
        "Deleting a saved run uses an in-page dialog that itemises the generated audio, submitted settings, metadata, history entry and Quick Phrases shortcut before any local data is removed. "
        "Long 32-character identifiers are truncated inside the fixed-width panel and remain available as secondary metadata, preventing horizontal overflow. "
        "Saved runs remain under user control and can be deleted permanently."
    ),
    "4.18": (
        "Figures 4.18 to 4.20 show three states of the nine-step Generate tour after a Quick Phrase has been saved. "
        "The tour dims the page, cuts a spotlight around the control under discussion and anchors a short explanatory card beside it. "
        "Figure 4.18 introduces navigation, Figure 4.19 explains the reference voice at step 5, and Figure 4.20 covers output format, watermarking and generation at step 7. "
        "The conditional Quick Phrases step is omitted when no saved phrase exists, so first-time users are not shown an empty feature."
    ),
    "4.21": (
        "Figure 4.21 shows the responsive Generate workflow at the tested 390 x 844 viewport. "
        "Quick Phrases are the first controls after the page heading, before model and reference configuration, because they are the fastest path for repeated assistive messages. "
        "A labelled phrase requests its stored audio directly and avoids model latency. "
        "Qwen, the shared FYP Demo voice and watermark state remain available through touch-sized controls. "
        "The primary action changes to Generating and is disabled while another job is active. "
        "The PWA uses the same local API and persistent records as the desktop interface."
    ),
    "4.22": (
        "Figure 4.22 presents mobile queue and history with All, Saved, Active, Done and Failed filters. "
        "Custom labels appear above the original script so a user can recognise an urgent phrase without losing its source text, and a star marks saved runs. "
        "Playing a completed run keeps the user on Jobs and opens the same lower player above navigation, with waveform seeking, download and a direct star control for saving or unsaving the clip. "
        "The detail sheet also exposes rename, restore and deletion actions; the deletion dialog itemises the local audio, metadata and Quick Phrases shortcut. "
        "The single-column cards avoid horizontal scrolling and the persistent local job metadata survives navigation and reloads."
    ),
    "4.23": (
        "The shared voice workflow is shown in Figure 4.23 with the privacy-safe FYP Demo record and its edit sheet. "
        "The user can preview, replace, rerecord, rename or correct the transcript. "
        "Voice deletion now uses the same in-page irreversible-action pattern as desktop: it lists the reference recording, transcript, preprocessing metadata and cached embeddings, and explains that existing generated clips remain. "
        "When saved phrases depend on the voice, the dialog warns that they continue to play but cannot be regenerated in that voice."
    ),
}


def media_hashes(path: Path) -> dict[str, str]:
    with zipfile.ZipFile(path) as archive:
        return {
            name: hashlib.sha256(archive.read(name)).hexdigest()
            for name in archive.namelist()
            if name.startswith("word/media/")
        }


def image_part_for_caption(doc: Document, caption_index: int) -> str:
    paragraphs = doc.paragraphs
    for index in range(caption_index - 1, max(-1, caption_index - 5), -1):
        blips = paragraphs[index]._p.xpath(".//a:blip")
        if not blips:
            continue
        relationship_id = blips[0].get(qn("r:embed"))
        part = paragraphs[index].part.related_parts[relationship_id]
        return str(part.partname).lstrip("/")
    raise RuntimeError(f"No image found before caption paragraph {caption_index}")


def replace_zip_entries(source: Path, destination: Path, replacements: dict[str, bytes]) -> None:
    with zipfile.ZipFile(source, "r") as incoming, zipfile.ZipFile(destination, "w") as outgoing:
        for info in incoming.infolist():
            outgoing.writestr(info, replacements.get(info.filename, incoming.read(info.filename)))


def main() -> None:
    missing = [filename for filename in FIGURES.values() if not (IMAGES / filename).exists()]
    if missing:
        raise RuntimeError(f"Missing refreshed screenshots: {missing}")

    BACKUPS.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y-%m-%d-%H%M%S")
    backup = BACKUPS / f"{DOCX.stem}-pre-ui-refresh-{stamp}.docx"
    shutil.copy2(DOCX, backup)
    before_hashes = media_hashes(DOCX)

    doc = Document(DOCX)
    caption_indexes: dict[str, int] = {}
    media_parts: dict[str, str] = {}
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph.style.name != "Caption":
            continue
        for number in FIGURES:
            if paragraph.text.startswith(f"Figure {number}:"):
                if number in caption_indexes:
                    raise RuntimeError(f"Duplicate body caption for Figure {number}")
                caption_indexes[number] = index
                media_parts[number] = image_part_for_caption(doc, index)

    if set(caption_indexes) != set(FIGURES):
        raise RuntimeError(f"Caption mismatch: {sorted(caption_indexes)}")
    if len(set(media_parts.values())) != len(FIGURES):
        raise RuntimeError(f"UI figures do not map to 13 unique media parts: {media_parts}")

    for number, text in CAPTIONS.items():
        doc.paragraphs[caption_indexes[number]].text = text
        doc.paragraphs[caption_indexes[number]].style = "Caption"
    for number, text in PROSE.items():
        caption_index = caption_indexes[number]
        for index in range(caption_index - 1, max(-1, caption_index - 5), -1):
            paragraph = doc.paragraphs[index]
            if paragraph.text.strip():
                paragraph.text = text
                paragraph.style = "Normal"
                break
        else:
            raise RuntimeError(f"No prose paragraph found before Figure {number}")

    uc10 = doc.tables[18].rows[4].cells[1]
    old = "3. The system asks the user to confirm the deletion."
    new = "3. The system itemises the reference recording, transcript, preprocessing metadata and cached embeddings, explains what generated clips remain, and asks the user to confirm permanent deletion."
    matches = [paragraph for paragraph in uc10.paragraphs if old in paragraph.text]
    already_updated = [paragraph for paragraph in uc10.paragraphs if new in paragraph.text]
    if len(matches) == 1:
        matches[0].text = matches[0].text.replace(old, new)
    elif len(matches) != 0 or len(already_updated) != 1:
        raise RuntimeError(
            f"Expected one old or updated UC-10 confirmation sentence, "
            f"found old={len(matches)} updated={len(already_updated)}"
        )

    prose_temp = DOCX.with_suffix(".ui-prose.tmp.docx")
    zip_temp = DOCX.with_suffix(".ui-refresh.tmp.docx")
    doc.save(prose_temp)
    replacements = {
        media_parts[number]: (IMAGES / filename).read_bytes()
        for number, filename in FIGURES.items()
    }
    replace_zip_entries(prose_temp, zip_temp, replacements)
    zip_temp.replace(DOCX)
    prose_temp.unlink(missing_ok=True)

    after_hashes = media_hashes(DOCX)
    changed = {name for name in set(before_hashes) | set(after_hashes) if before_hashes.get(name) != after_hashes.get(name)}
    expected = set(replacements)
    required_changed = {media_parts[number] for number in ("4.12", "4.18", "4.19", "4.20", "4.21", "4.22")}
    if not changed.issubset(expected) or not required_changed.issubset(changed):
        shutil.copy2(backup, DOCX)
        raise RuntimeError(
            f"Unexpected media changes; restored backup. Allowed {sorted(expected)}, "
            f"required {sorted(required_changed)}, changed {sorted(changed)}"
        )

    print(f"Updated {DOCX}")
    print(f"Backup: {backup}")
    print("Changed media: " + ", ".join(sorted(changed)))
    unchanged = expected - changed
    if unchanged:
        print("Byte-identical replacements: " + ", ".join(sorted(unchanged)))


if __name__ == "__main__":
    main()
